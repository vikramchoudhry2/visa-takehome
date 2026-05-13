"""Parse a Vertex client briefing .docx into a `Brief` object.

Strategy:
1. Walk the document body in document order, interleaving paragraphs and
   tables (python-docx exposes them via the underlying XML element tree).
2. Identify section boundaries by fuzzy-matching paragraph text against
   each body header title (see `BODY_HEADER_IDS_IN_ORDER` and
   `_body_header_display_title`).
3. Bucket subsequent paragraphs and tables into the active section until
   the next header is found.
3b. Some Word templates place sections 03–07 (and beyond) in one 2-column
   table (``<label> | <value>``). The parser splits such tables into
   synthetic per-section rows. **Messy real files** often merge the Key
   Facts umbrella label with long body copy in the same cell, use
   "President" instead of "executive", or put two attendee questions in one
   cell; we match headers on the **first line** / **case-insensitive prefix**
   of the left cell and skip non-matching rows instead of rejecting the
   whole table.
4. Special-case section 01 (header icon) and section 02 (title) — these
   live in the header / first paragraph respectively, not as their own
   labeled body sections.
5. Always emit all 15 sections in template order; sections we couldn't
   find are emitted with `present=False`.
6. The Word template exposes one body header, "What are the key facts
   about the client?"; that text is the **umbrella section title** in the
   Word file, not a separate row in the review output table. The parser
   buckets that block under an internal parse id, then splits it into
   four review-table rows (8A–8D). Combined Key Facts tables are attached
   to the 8A section so structure checkers can still read label cells.
   Subsections may appear as ``Label: body`` paragraphs in one cell, or as
   rows in a 2-column table (label cell | value cell), including a blank
   label row when Vertex Overview copy sits in the value column only.

Fuzzy matching uses rapidfuzz with a conservative score threshold so
small wording deltas (e.g., en-dash vs hyphen, smart quotes) don't break
section detection.
"""

from __future__ import annotations

import functools
import io
import re
from collections.abc import Iterator
from dataclasses import dataclass

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from rapidfuzz import fuzz

from core.models import (
    SECTION_TEMPLATE,
    Brief,
    ImageRef,
    ParsedTable,
    Section,
)

# Body header "What are the key facts…" is parsed into this internal bucket
# (not a template row id), then split into rows 08a–08d for the review table.
KF_BODY_PARSE_ID = "__kf_body__"
KEY_FACTS_HEADER_DISPLAY = "What are the key facts about the client?"

FUZZY_HEADER_THRESHOLD = 82

BODY_HEADER_IDS_IN_ORDER: tuple[str, ...] = (
    "03_client_name_type",
    "04_meeting_objective",
    "05_client_markets",
    "06_client_share",
    "07_current_business",
    KF_BODY_PARSE_ID,
    "09_exec_messages",
    "10_client_topics",
    "11_meeting_with",
    "12_vertex_attendees",
)

# Notable Changes is a subsection inside Key Facts, not a standalone body header.
_NOTABLE_CHANGES_BODY_RE = re.compile(
    r"notable\s+changes\s*:\s*(.*)$",
    re.IGNORECASE | re.DOTALL,
)

@dataclass
class _Block:
    """A paragraph or table from the body, in document order."""

    kind: str
    paragraph: Paragraph | None = None
    table: Table | None = None
    # Synthetic 2-col row extracted from a stacked template table (label | value).
    synthetic_value: tuple[str, str] | None = None


def parse_brief(source: bytes | io.BytesIO | str) -> Brief:
    """Parse a docx file (path, bytes, or stream) into a `Brief`."""
    if isinstance(source, bytes):
        source = io.BytesIO(source)
    doc = Document(source)
    return _parse_document(doc)


def _parse_document(doc: DocxDocument) -> Brief:
    blocks = list(_iter_body_blocks(doc))
    header_images = _extract_header_images(doc)
    title_text, title_idx = _extract_title(blocks)
    header_plain = _collect_header_paragraph_text(doc)
    body_plain = _collect_full_text(blocks, doc)
    full_text = "\n".join(p for p in (header_plain, body_plain) if p.strip())

    body_section_blocks = _split_body_into_sections(
        blocks, start_idx=title_idx + 1 if title_idx is not None else 0
    )

    kf_blocks = body_section_blocks.get(KF_BODY_PARSE_ID)
    key_facts_combined: Section | None = None
    if kf_blocks:
        key_facts_combined = _section_from_blocks(
            KF_BODY_PARSE_ID,
            KEY_FACTS_HEADER_DISPLAY,
            0,
            kf_blocks,
        )

    sections: list[Section] = []
    for order, (sid, display_title) in enumerate(SECTION_TEMPLATE, start=1):
        if sid == "01_header_icon":
            sections.append(
                Section(
                    id=sid,
                    title=display_title,
                    order=order,
                    present=bool(header_images),
                    images=tuple(header_images),
                )
            )
        elif sid == "02_title":
            sections.append(
                Section(
                    id=sid,
                    title=display_title,
                    order=order,
                    present=title_text is not None,
                    raw_text=title_text or "",
                )
            )
        elif sid == "08a_business_overview":
            sections.append(
                _key_facts_subsection(
                    key_facts_combined,
                    sid,
                    display_title,
                    order,
                    category="business overview",
                    key_facts_tables=(
                        key_facts_combined.tables
                        if key_facts_combined is not None
                        else ()
                    ),
                )
            )
        elif sid == "08b_competition":
            sections.append(
                _key_facts_subsection(
                    key_facts_combined,
                    sid,
                    display_title,
                    order,
                    category="competition",
                )
            )
        elif sid == "08c_vertex_overview":
            sections.append(
                _key_facts_subsection(
                    key_facts_combined,
                    sid,
                    display_title,
                    order,
                    category="vertex overview",
                )
            )
        elif sid == "08d_notable_changes":
            sections.append(
                _key_facts_subsection_notable(
                    key_facts_combined, sid, display_title, order
                )
            )
        else:
            blocks_for_section = body_section_blocks.get(sid)
            if not blocks_for_section:
                sections.append(
                    Section(id=sid, title=display_title, order=order, present=False)
                )
            else:
                sections.append(
                    _section_from_blocks(sid, display_title, order, blocks_for_section)
                )

    client_name = _infer_client_name(title_text)

    return Brief(
        client_name=client_name,
        title_text=title_text,
        header_images=tuple(header_images),
        sections=tuple(sections),
        full_text=full_text,
        key_facts_combined_raw=(
            key_facts_combined.raw_text if key_facts_combined is not None else ""
        ),
    )


def _iter_body_blocks(doc: DocxDocument) -> Iterator[_Block]:
    """Yield paragraphs and tables in document order."""
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            yield _Block(kind="paragraph", paragraph=Paragraph(child, doc))
        elif tag == qn("w:tbl"):
            yield _Block(kind="table", table=Table(child, doc))


def _extract_header_images(doc: DocxDocument) -> list[ImageRef]:
    images: list[ImageRef] = []
    for sect in doc.sections:
        for header in (sect.header, sect.first_page_header, sect.even_page_header):
            if header is None:
                continue
            for shape in _iter_inline_shapes(header._element):
                images.append(
                    ImageRef(
                        location="header",
                        width_emu=shape.get("cx"),
                        height_emu=shape.get("cy"),
                    )
                )
    return images


def _iter_inline_shapes(element) -> Iterator[dict]:
    """Yield {cx, cy} dicts for every embedded image under `element`."""
    for blip in element.iter(qn("a:blip")):
        extent = None
        for parent in blip.iterancestors():
            extent = parent.find(qn("wp:extent"))
            if extent is not None:
                break
        cx = int(extent.get("cx")) if extent is not None and extent.get("cx") else None
        cy = int(extent.get("cy")) if extent is not None and extent.get("cy") else None
        yield {"cx": cx, "cy": cy}


def _has_inline_image(element) -> bool:
    return next(iter(element.iter(qn("a:blip"))), None) is not None


def _extract_title(blocks: list[_Block]) -> tuple[str | None, int | None]:
    """Find the brief title. By spec, must be `[Client Name] Meeting Brief`.

    Search the first ~10 non-empty paragraphs for one containing
    'Meeting Brief'.
    """
    seen = 0
    for idx, block in enumerate(blocks):
        if block.kind != "paragraph":
            continue
        text = block.paragraph.text.strip()
        if not text:
            continue
        seen += 1
        if "meeting brief" in text.lower():
            return text, idx
        if seen >= 10:
            break
    return None, None


def _collect_header_paragraph_text(doc: DocxDocument) -> str:
    """Plain text from the first section header (for spelling/formatting on header copy)."""
    parts: list[str] = []
    for section in doc.sections:
        for p in section.header.paragraphs:
            if p.text.strip():
                parts.append(p.text)
    return "\n".join(parts)


def _collect_full_text(blocks: list[_Block], doc: DocxDocument) -> str:
    parts: list[str] = []
    for block in blocks:
        if block.kind == "paragraph":
            txt = block.paragraph.text
            if txt:
                parts.append(txt)
        else:
            for row in block.table.rows:
                for cell in row.cells:
                    cell_text = "\n".join(p.text for p in cell.paragraphs if p.text)
                    if cell_text:
                        parts.append(cell_text)
    return "\n".join(parts)


def _cell_plain_text(cell: _Cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs if p.text).strip()


def _body_header_display_title(section_id: str) -> str:
    """Display string used for fuzzy-matching a body section header."""
    if section_id == KF_BODY_PARSE_ID:
        return KEY_FACTS_HEADER_DISPLAY
    for sid, title in SECTION_TEMPLATE:
        if sid == section_id:
            return title
    raise KeyError(section_id)


def _best_header_fuzzy(text: str) -> tuple[str | None, int]:
    """Return (best matching section id, score) for fuzzy header matching."""
    stripped = text.strip()
    if not stripped:
        return None, 0
    normalized = _normalize(stripped)
    best_id: str | None = None
    best_score = 0
    for sid in BODY_HEADER_IDS_IN_ORDER:
        target = _normalize(_body_header_display_title(sid))
        score = fuzz.ratio(normalized, target)
        if score > best_score:
            best_score = score
            best_id = sid
    return best_id, best_score


@functools.cache
def _body_header_specs_longest_first() -> tuple[tuple[str, str], ...]:
    """Template (section_id, display title) pairs, longest title first for prefix checks."""
    return tuple(
        sorted(
            ((sid, _body_header_display_title(sid)) for sid in BODY_HEADER_IDS_IN_ORDER),
            key=lambda x: len(x[1]),
            reverse=True,
        )
    )


def _line_header_match(line: str) -> tuple[str | None, int]:
    """Match a single line or short span to a body header (prefix, then fuzzy)."""
    line_st = line.strip()
    if not line_st:
        return None, 0
    lower = line_st.lower()
    for sid, display in _body_header_specs_longest_first():
        dlow = display.lower()
        if lower.startswith(dlow):
            return sid, 100
    return _best_header_fuzzy(line_st)


def _remainder_after_display_prefix(line: str, display: str) -> str:
    """Strip a leading template display title (case-insensitive) from ``line``."""
    s, d = line.strip(), display.strip()
    if s.lower().startswith(d.lower()):
        return s[len(d) :].strip()
    return s


def _dedupe_value_right(left: str, right: str) -> str:
    r = right.strip()
    if not r or r == left.strip():
        return ""
    return r


_COMPOUND_HEADER_LINE_MAX_CHARS = 200
_MIN_CHARS_BEFORE_EMBEDDED_CLIENT_TOPICS = 8


def _try_compound_two_line_headers(left: str, right: str) -> list[tuple[str, str, str]]:
    """Two stacked questions in one cell (e.g. Who am I… + Who is joining…)."""
    lines = [ln.strip() for ln in left.splitlines() if ln.strip()]
    if len(lines) < 2:
        return []
    l0, l1 = lines[0], lines[1]
    if len(l0) > _COMPOUND_HEADER_LINE_MAX_CHARS:
        return []
    if len(l1) > _COMPOUND_HEADER_LINE_MAX_CHARS:
        return []
    m0_sid, m0_sc = _line_header_match(l0)
    m1_sid, m1_sc = _line_header_match(l1)
    if (
        m0_sid is None
        or m1_sid is None
        or m0_sid == m1_sid
        or m0_sc < FUZZY_HEADER_THRESHOLD
        or m1_sc < FUZZY_HEADER_THRESHOLD
    ):
        return []
    disp1 = _body_header_display_title(m1_sid)
    v1 = _remainder_after_display_prefix(l1, disp1)
    tail = "\n".join(lines[2:]).strip()
    vr = _dedupe_value_right(left, right)
    val_second = "\n".join(x for x in (v1, tail, vr) if x)
    return [(m0_sid, l0, ""), (m1_sid, disp1, val_second)]


@dataclass(frozen=True)
class _LeftCellHeaderPick:
    section_id: str
    score: int
    label: str
    remainder: str


def _pick_header_from_left_cell(left: str) -> _LeftCellHeaderPick | None:
    """Prefer a header on the first line when the cell also contains long body copy."""
    stripped = left.strip()
    if not stripped:
        return None
    parts = stripped.split("\n", 1)
    first = parts[0].strip()
    after_first = parts[1].strip() if len(parts) > 1 else ""
    candidates: list[tuple[str, str]] = [(first, after_first)]
    if stripped != first:
        candidates.append((stripped, ""))
    best: _LeftCellHeaderPick | None = None
    for label, remainder in candidates:
        sid, score = _line_header_match(label)
        if sid is None or score == 0:
            continue
        cand = _LeftCellHeaderPick(
            section_id=sid, score=score, label=label, remainder=remainder
        )
        if best is None or score > best.score or (
            score == best.score and len(label) < len(best.label)
        ):
            best = cand
    return best


def _physical_row_to_stack_entries(left: str, right: str) -> list[tuple[str, str, str]]:
    compound = _try_compound_two_line_headers(left, right)
    if compound:
        return compound
    picked = _pick_header_from_left_cell(left)
    if picked is None or picked.score < FUZZY_HEADER_THRESHOLD:
        return []
    vr = _dedupe_value_right(left, right)
    body = "\n".join(x for x in (picked.remainder, vr) if x)
    return _split_embedded_template_headers(picked.section_id, picked.label, body)


def _split_embedded_template_headers(
    primary_sid: str, primary_label: str, combined: str
) -> list[tuple[str, str, str]]:
    """When one cell holds executive messages then client topics (pipe-separated), split."""
    if primary_sid != "09_exec_messages" or not combined.strip():
        return [(primary_sid, primary_label, combined)]
    disp10 = _body_header_display_title("10_client_topics")
    m = re.search(re.escape(disp10), combined, flags=re.IGNORECASE)
    if m is None or m.start() < _MIN_CHARS_BEFORE_EMBEDDED_CLIENT_TOPICS:
        return [(primary_sid, primary_label, combined)]
    exec_body = combined[: m.start()].strip()
    topics_body = combined[m.end() :].strip()
    if not topics_body:
        return [(primary_sid, primary_label, combined)]
    out: list[tuple[str, str, str]] = [(primary_sid, primary_label, exec_body)]
    out.append(("10_client_topics", disp10, topics_body))
    return out


def _merge_stack_rows_by_section(rows: list[tuple[str, str, str]]) -> list[tuple[str, str, str]]:
    """Merge duplicate section ids (multiple physical rows) preserving first-seen order."""
    order: list[str] = []
    acc: dict[str, tuple[str, str, str]] = {}
    for sid, lab, val in rows:
        if sid not in acc:
            order.append(sid)
            acc[sid] = (sid, lab, val)
            continue
        _prev_sid, prev_lab, prev_val = acc[sid]
        merged_val = "\n".join(x for x in (prev_val, val) if x.strip())
        acc[sid] = (sid, prev_lab, merged_val)
    return [acc[s] for s in order]


def _classify_nested_table(table: Table) -> str | None:
    """Best-effort: which section_id does this nested table belong to?

    Looks at the header row. Returns ``None`` when no confident match —
    callers fall back to the last stack entry from the same physical row.
    """
    if not table.rows:
        return None
    first_row = " | ".join(c.text.strip().lower() for c in table.rows[0].cells)
    if (
        "name/titles" in first_row
        or "previously met" in first_row
        or ("photo" in first_row and "bio" in first_row)
    ):
        return "11_meeting_with"
    return None


def _stacked_two_col_template_table(
    table: Table,
) -> list[tuple[str, str, str, tuple[Table, ...]]] | None:
    """Split a 2-column table into synthetic ``(section_id, label, value, nested_tables)`` rows.

    Some templates put the entire brief inside one outer 2-column table
    where the right cell can hold *nested* tables (e.g. the 4-column
    attendee table for section 11). We surface those nested tables on
    each stack entry so structure checkers can inspect them as real
    `ParsedTable` objects rather than flattened text.

    Tolerates label+body merged in the left cell, optional ``President`` vs
    ``executive`` wording, two attendee questions in one cell, and non-matching
    rows (skipped) once at least one header row matched.
    """
    rows_out: list[tuple[str, str, str, tuple[Table, ...]]] = []
    for row in table.rows:
        if len(row.cells) < 2:
            return None
        left = _cell_plain_text(row.cells[0])
        right_cell = row.cells[1]
        right = _cell_plain_text(right_cell)
        entries = _physical_row_to_stack_entries(left, right)
        if not entries:
            continue
        nested = tuple(right_cell.tables)
        # Route each nested table to a specific section when its header
        # content classifies cleanly; otherwise default to the last
        # entry from this physical row.
        entry_sids = [sid for sid, _, _ in entries]
        per_entry_nested: dict[int, list[Table]] = {i: [] for i in range(len(entries))}
        for nt in nested:
            classified = _classify_nested_table(nt)
            if classified is not None and classified in entry_sids:
                target_idx = entry_sids.index(classified)
            else:
                target_idx = len(entries) - 1
            per_entry_nested[target_idx].append(nt)
        for idx, (sid, lab, val) in enumerate(entries):
            attach = tuple(per_entry_nested[idx])
            rows_out.append((sid, lab, val, attach))
    if not rows_out:
        return None
    return _merge_stack_rows_with_nested(rows_out)


def _merge_stack_rows_with_nested(
    rows: list[tuple[str, str, str, tuple[Table, ...]]],
) -> list[tuple[str, str, str, tuple[Table, ...]]]:
    """Same as `_merge_stack_rows_by_section` but preserves nested tables."""
    order: list[str] = []
    acc: dict[str, tuple[str, str, str, tuple[Table, ...]]] = {}
    for sid, lab, val, nested in rows:
        if sid not in acc:
            order.append(sid)
            acc[sid] = (sid, lab, val, nested)
            continue
        _prev_sid, prev_lab, prev_val, prev_nested = acc[sid]
        merged_val = "\n".join(x for x in (prev_val, val) if x.strip())
        acc[sid] = (sid, prev_lab, merged_val, prev_nested + nested)
    return [acc[s] for s in order]


def _split_body_into_sections(
    blocks: list[_Block], start_idx: int
) -> dict[str, list[_Block]]:
    """Walk body blocks, splitting them into per-section block lists."""
    section_buckets: dict[str, list[_Block]] = {}
    current_id: str | None = None

    for block in blocks[start_idx:]:
        if block.kind == "table":
            stacked = _stacked_two_col_template_table(block.table)
            if stacked is not None:
                for sid, left, right, nested in stacked:
                    bucket = section_buckets.setdefault(sid, [])
                    bucket.append(
                        _Block(
                            kind="synthetic_value",
                            synthetic_value=(left, right),
                        )
                    )
                    for nt in nested:
                        bucket.append(_Block(kind="table", table=nt))
                current_id = stacked[-1][0]
                continue

        matched_id = _match_block_to_header(block)
        if matched_id is not None and matched_id != current_id:
            current_id = matched_id
            section_buckets.setdefault(current_id, [])
            continue
        if current_id is None:
            continue
        section_buckets.setdefault(current_id, []).append(block)

    return section_buckets


def _match_block_to_header(block: _Block) -> str | None:
    """Return the section id this block is a header for, or None.

    Matches if the block is a paragraph whose text closely resembles a
    template header. Tables are not paragraph headers; stacked 2-column
    tables whose first column repeats template titles are handled in
    `_split_body_into_sections` via `_stacked_two_col_template_table`.
    """
    if block.kind != "paragraph" or block.paragraph is None:
        return None
    text = block.paragraph.text.strip()
    if not text:
        return None
    best_id, best_score = _best_header_fuzzy(text)
    if best_score >= FUZZY_HEADER_THRESHOLD:
        return best_id
    return None


_NORMALIZE_RE = re.compile(r"[\s\u2010-\u2015\-]+")


def _normalize(text: str) -> str:
    text = text.replace("\u2019", "'").replace("\u2018", "'")
    text = text.replace("\u201c", '"').replace("\u201d", '"')
    text = _NORMALIZE_RE.sub(" ", text)
    return text.strip().lower()


def _section_from_blocks(
    section_id: str, title: str, order: int, blocks: list[_Block]
) -> Section:
    text_parts: list[str] = []
    bullets: list[str] = []
    tables: list[ParsedTable] = []
    images: list[ImageRef] = []

    for block in blocks:
        if block.kind == "synthetic_value":
            if block.synthetic_value is not None:
                label, value = block.synthetic_value
                text_parts.append(label)
                text_parts.append(value)
                tables.append(ParsedTable(rows=((label, value),), has_images=False))
            continue
        if block.kind == "paragraph":
            p = block.paragraph
            txt = p.text
            if txt.strip():
                text_parts.append(txt)
                if _is_list_paragraph(p):
                    bullets.append(txt.strip())
            for _shape in _iter_inline_shapes(p._element):
                images.append(ImageRef(location="body"))
        elif block.kind == "table" and block.table is not None:
            tbl = block.table
            tables.append(_parse_table(tbl))
            for cell in _iter_table_cells(tbl):
                cell_text = "\n".join(par.text for par in cell.paragraphs if par.text)
                if cell_text.strip():
                    text_parts.append(cell_text)
                for par in cell.paragraphs:
                    if _is_list_paragraph(par) and par.text.strip():
                        bullets.append(par.text.strip())
                    if _has_inline_image(par._element):
                        images.append(ImageRef(location="table_cell"))

    return Section(
        id=section_id,
        title=title,
        order=order,
        present=True,
        raw_text="\n".join(text_parts),
        bullets=tuple(bullets),
        tables=tuple(tables),
        images=tuple(images),
    )


def _cell_text_in_document_order(cell: _Cell) -> str:
    """Plain text from a table cell, including paragraphs inside nested tables.

    Word often wraps body copy in a nested ``<w:tbl>`` inside a cell. The
    high-level ``cell.paragraphs`` API only lists direct child paragraphs,
    so those cells would otherwise parse as empty and semantic checks
    falsely report an empty Bio column.
    """
    parts: list[str] = []
    for child in cell._tc.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            p = Paragraph(child, cell)
            if p.text.strip():
                parts.append(p.text)
        elif tag == qn("w:tbl"):
            inner = Table(child, cell)
            for inner_row in inner.rows:
                for inner_cell in inner_row.cells:
                    nested = _cell_text_in_document_order(inner_cell)
                    if nested.strip():
                        parts.append(nested)
    return "\n".join(parts).strip()


def _cell_has_inline_image_deep(cell: _Cell) -> bool:
    """True if this cell or any nested table cell contains an inline image."""
    for child in cell._tc.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            p = Paragraph(child, cell)
            if _has_inline_image(p._element):
                return True
        elif tag == qn("w:tbl"):
            inner = Table(child, cell)
            for inner_row in inner.rows:
                for inner_cell in inner_row.cells:
                    if _cell_has_inline_image_deep(inner_cell):
                        return True
    return False


def _iter_table_cells(table: Table) -> Iterator[_Cell]:
    """Yield each logical table cell once, row-major.

    Deduplicate by python ``Cell`` object identity (not ``cell._tc``): in
    some merged layouts python-docx surfaces distinct ``Cell`` instances
    that share the same underlying ``_tc`` element, and deduping by
    ``_tc`` would drop real rows (e.g. Key Facts label/value pairs).
    """
    seen: set[int] = set()
    for row in table.rows:
        for cell in row.cells:
            cid = id(cell)
            if cid in seen:
                continue
            seen.add(cid)
            yield cell


def _parse_table(table: Table) -> ParsedTable:
    rows: list[tuple[str, ...]] = []
    cell_image_rows: list[tuple[bool, ...]] = []
    has_images = False
    for row in table.rows:
        cells: list[str] = []
        cell_imgs: list[bool] = []
        for cell in row.cells:
            cell_text = _cell_text_in_document_order(cell)
            cells.append(cell_text)
            cell_has_img = _cell_has_inline_image_deep(cell)
            cell_imgs.append(cell_has_img)
            if cell_has_img:
                has_images = True
        rows.append(tuple(cells))
        cell_image_rows.append(tuple(cell_imgs))
    return ParsedTable(
        rows=tuple(rows),
        has_images=has_images,
        cell_images=tuple(cell_image_rows),
    )


def _is_list_paragraph(paragraph: Paragraph) -> bool:
    """A paragraph is a list/bullet if it has a numPr element OR a list-style style.

    Word stores list membership in two places: a `<w:numPr>` element on
    the paragraph (or inherited from its style), or via a style whose
    name begins with `List`. We treat either as a bullet.
    """
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is not None and pPr.find(qn("w:numPr")) is not None:
        return True
    style = paragraph.style
    if style is None or style.name is None:
        return False
    name = style.name.lower()
    return name.startswith("list ") or name == "list bullet" or name == "list number"


def _normalize_key_facts_line_heading(head: str) -> str:
    """Strip bullets, '8A.', 'Only if applicable:', etc. before label compare."""
    s = head.strip()
    s = re.sub(r"^\s*\|+\s*", "", s)
    s = re.sub(r"^[\u2022•\-\*\s]+", "", s)
    s = re.sub(r"^only if applicable:\s*", "", s, flags=re.IGNORECASE)
    s = re.sub(r"^\d{1,2}[a-d]\s*[\.\)]\s*", "", s, flags=re.IGNORECASE)
    return s.strip().lower()


def _label_cell_to_key_facts_category(left: str) -> str | None:
    """Map a Key Facts table label cell to canonical subsection id text."""
    if not left or not left.strip():
        return None
    head = left.strip().split(":", 1)[0]
    n = _normalize_key_facts_line_heading(head)
    n = re.sub(r"\s*\(if applicable\)\s*$", "", n, flags=re.IGNORECASE).strip()
    if n in ("business overview", "competition", "vertex overview", "notable changes"):
        return n
    return None


def _iter_key_facts_label_value_pairs(section: Section) -> list[tuple[str, str]]:
    """(label cell, value cell) for each row of every 2-column table in the section."""
    pairs: list[tuple[str, str]] = []
    for tbl in section.tables:
        if tbl.num_cols < 2:
            continue
        for row in tbl.rows:
            if len(row) < 2:
                continue
            pairs.append((row[0].strip(), row[1].strip()))
    return pairs


def _combined_key_facts_has_content(combined: Section) -> bool:
    if combined.raw_text.strip():
        return True
    for tbl in combined.tables:
        for row in tbl.rows:
            if any(cell.strip() for cell in row):
                return True
    return False


def _value_reads_as_vertex_overview(right: str) -> bool:
    """Heuristic for Vertex copy when the label cell is left blank (common in templates)."""
    t = right.lower()
    if re.search(r"cards?\s+in\s+force", t):
        return True
    if "portfolio" in t and re.search(r"\d", t):
        return True
    if "vertex" in t and ("card" in t or "portfolio" in t or "pv" in t):
        return True
    return False


def _extract_key_facts_category_from_tables(section: Section, category: str) -> str | None:
    """Match Key Facts laid out as a 2-column Word table (label | body per row)."""
    pairs = _iter_key_facts_label_value_pairs(section)
    for left, right in pairs:
        if not right.strip():
            continue
        lbl = _label_cell_to_key_facts_category(left)
        if lbl == category:
            return right.strip()
    if category == "vertex overview":
        for left, right in pairs:
            if left.strip() or not right.strip():
                continue
            if _value_reads_as_vertex_overview(right):
                return right.strip()
        for left, right in pairs:
            if left.strip() or not right.strip():
                continue
            return right.strip()
    return None


def _extract_key_facts_category_body_inline(raw: str, category: str) -> str | None:
    """Return body for ``Label: body`` lines inside one cell / paragraph block."""
    for line in (x.strip() for x in raw.split("\n") if x.strip()):
        if ":" not in line:
            continue
        head, rest = line.split(":", 1)
        norm = _normalize_key_facts_line_heading(head)
        if "notable changes" in norm:
            continue
        if norm == category:
            return rest.strip()
    return None


def _extract_key_facts_subsection_body(combined: Section, category: str) -> str | None:
    """Prefer 2-column Key Facts table rows; fall back to inline ``Label:`` lines."""
    from_tbl = _extract_key_facts_category_from_tables(combined, category)
    if from_tbl is not None:
        return from_tbl
    return _extract_key_facts_category_body_inline(combined.raw_text, category)


def _key_facts_subsection(
    combined: Section | None,
    section_id: str,
    title: str,
    order: int,
    *,
    category: str,
    key_facts_tables: tuple[ParsedTable, ...] = (),
) -> Section:
    if combined is None or not _combined_key_facts_has_content(combined):
        return Section(id=section_id, title=title, order=order, present=False)
    body = _extract_key_facts_subsection_body(combined, category)
    if not body or not body.strip():
        return Section(id=section_id, title=title, order=order, present=False)
    tables_out: tuple[ParsedTable, ...] = ()
    if category == "business overview" and key_facts_tables:
        tables_out = key_facts_tables
    return Section(
        id=section_id,
        title=title,
        order=order,
        present=True,
        raw_text=body,
        tables=tables_out,
    )


def _key_facts_subsection_notable(
    combined: Section | None, section_id: str, title: str, order: int
) -> Section:
    """Notable Changes is optional: always `present=True` so it is never
    flagged as a missing section; empty body is normal."""
    if combined is None or not _combined_key_facts_has_content(combined):
        return Section(
            id=section_id,
            title=title,
            order=order,
            present=True,
            raw_text="",
        )
    body = _extract_notable_changes_from_combined(combined)
    return Section(
        id=section_id,
        title=title,
        order=order,
        present=True,
        raw_text=body,
    )


def _extract_notable_changes_from_combined(section: Section) -> str:
    from_tbl = _extract_key_facts_category_from_tables(section, "notable changes")
    if from_tbl is not None:
        return from_tbl
    return _extract_notable_changes_body_inline(section.raw_text)


def _extract_notable_changes_body_inline(key_facts_raw: str) -> str:
    """Return text after the 'Notable Changes:' marker, or empty if absent."""
    for paragraph in (p.strip() for p in key_facts_raw.split("\n") if p.strip()):
        if "notable changes" not in paragraph.lower():
            continue
        match = _NOTABLE_CHANGES_BODY_RE.search(paragraph)
        if match:
            return match.group(1).strip()
    return ""


def _infer_client_name(title_text: str | None) -> str | None:
    if not title_text:
        return None
    match = re.match(r"^(.*?)\s+Meeting Brief\b", title_text.strip(), re.IGNORECASE)
    if not match:
        return None
    name = match.group(1).strip().strip("[]")
    return name or None


