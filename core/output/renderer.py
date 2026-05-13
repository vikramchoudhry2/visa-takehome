"""Render a `ReviewReport` to markdown, pandas DataFrame, .docx, and CSV.

The required output format from the assessment is a 3-column table with
headers:
  Brief section | Formatting and Grammar Feedback | Brief Section Feedback

Columns 2 and 3 list feedback as markdown bullets: each item is a line
starting with ``- ``, and consecutive bullets are separated by a single
newline. Missing sections are still listed and flagged. Compliant sections
show ``- No issues found``.
"""

from __future__ import annotations

import csv
import io
from typing import TYPE_CHECKING

from docx import Document
from docx.shared import Inches

from core.models import ReviewReport, SectionReport

if TYPE_CHECKING:
    import pandas as pd

NO_ISSUES = "No issues found"
SECTION_MISSING = "Section missing - please add this section to the brief"
COLUMN_HEADERS = ("Brief section", "Formatting and Grammar Feedback", "Brief Section Feedback")


def to_dataframe(report: ReviewReport) -> "pd.DataFrame":
    import pandas as pd

    rows = []
    for sr in report.rows:
        rows.append(
            {
                COLUMN_HEADERS[0]: sr.title,
                COLUMN_HEADERS[1]: _format_bullets(sr.formatting_findings, missing=False),
                COLUMN_HEADERS[2]: _format_bullets(sr.section_findings, missing=not sr.present),
            }
        )
    return pd.DataFrame(rows, columns=list(COLUMN_HEADERS))


def to_markdown(report: ReviewReport) -> str:
    out = ["| " + " | ".join(COLUMN_HEADERS) + " |", "|" + "|".join(["---"] * 3) + "|"]
    for sr in report.rows:
        col1 = sr.title.replace("|", "\\|")
        col2 = _format_bullets_md(sr.formatting_findings, missing=False)
        col3 = _format_bullets_md(sr.section_findings, missing=not sr.present)
        out.append(f"| {col1} | {col2} | {col3} |")
    return "\n".join(out)


def to_csv(report: ReviewReport) -> bytes:
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(COLUMN_HEADERS)
    for sr in report.rows:
        writer.writerow(
            [
                sr.title,
                _format_bullets(sr.formatting_findings, missing=False),
                _format_bullets(sr.section_findings, missing=not sr.present),
            ]
        )
    return buf.getvalue().encode("utf-8")


def to_docx(report: ReviewReport) -> bytes:
    doc = Document()
    title = report.client_name or "Vertex"
    doc.add_heading(f"{title} - Brief Review", level=1)
    table = doc.add_table(rows=1 + len(report.rows), cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(COLUMN_HEADERS):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for r_idx, sr in enumerate(report.rows, start=1):
        row = table.rows[r_idx].cells
        row[0].text = sr.title
        _write_bullet_cell(row[1], sr.formatting_findings, missing=False)
        _write_bullet_cell(row[2], sr.section_findings, missing=not sr.present)
    for col_idx, width_in in enumerate((1.7, 2.6, 2.6)):
        for row in table.rows:
            row.cells[col_idx].width = Inches(width_in)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _format_bullets(findings: tuple, missing: bool) -> str:
    """One markdown-style bullet per line; consecutive bullets separated by a single newline."""
    if missing and not findings:
        return f"- {SECTION_MISSING}"
    if not findings:
        return f"- {NO_ISSUES}"
    return "\n".join(f"- {_finding_text(f)}" for f in findings)


def _format_bullets_md(findings: tuple, missing: bool) -> str:
    """Same bullet layout as `_format_bullets`, with pipe characters escaped for markdown tables."""
    if missing and not findings:
        return f"- {SECTION_MISSING.replace('|', '\\|')}"
    if not findings:
        return f"- {NO_ISSUES.replace('|', '\\|')}"
    return "\n".join(f"- {_finding_text(f).replace('|', '\\|')}" for f in findings)


def _finding_text(finding) -> str:
    return finding.message


def _write_bullet_cell(cell, findings: tuple, missing: bool) -> None:
    cell.text = ""
    paragraphs = cell.paragraphs
    if missing and not findings:
        paragraphs[0].text = SECTION_MISSING
        paragraphs[0].style = cell.part.document.styles["List Bullet"]
        return
    if not findings:
        paragraphs[0].text = NO_ISSUES
        paragraphs[0].style = cell.part.document.styles["List Bullet"]
        return
    first = True
    for f in findings:
        if first:
            paragraphs[0].text = _finding_text(f)
            paragraphs[0].style = cell.part.document.styles["List Bullet"]
            first = False
        else:
            p = cell.add_paragraph(_finding_text(f), style="List Bullet")
            _ = p


def summary_counts(report: ReviewReport) -> dict[str, int]:
    formatting = sum(len(sr.formatting_findings) for sr in report.rows)
    section = sum(len(sr.section_findings) for sr in report.rows)
    missing = sum(1 for sr in report.rows if not sr.present)
    sections_clean = sum(
        1
        for sr in report.rows
        if sr.present and not sr.formatting_findings and not sr.section_findings
    )
    return {
        "formatting": formatting,
        "section": section,
        "missing_sections": missing,
        "clean_sections": sections_clean,
        "total_sections": len(report.rows),
    }


_ = SectionReport  # re-export for type hinting
