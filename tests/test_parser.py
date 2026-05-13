"""Parser tests against the synthetic fixtures."""

from __future__ import annotations

from core.docx_parser import parse_brief
from core.models import SECTION_IDS, Brief


def test_parses_all_15_sections(clean_brief: Brief) -> None:
    assert len(clean_brief.sections) == 15
    assert tuple(s.id for s in clean_brief.sections) == SECTION_IDS


def test_sections_are_present_in_clean_brief(clean_brief: Brief) -> None:
    assert all(s.present for s in clean_brief.sections)


def test_extracts_client_name_and_title(clean_brief: Brief) -> None:
    assert clean_brief.client_name == "Acme Bank"
    assert clean_brief.title_text == "Acme Bank Meeting Brief"


def test_header_image_detected(clean_brief: Brief) -> None:
    assert len(clean_brief.header_images) == 1
    assert clean_brief.header_images[0].location == "header"


def test_two_column_tables_have_two_columns(clean_brief: Brief) -> None:
    for sid in (
        "03_client_name_type",
        "04_meeting_objective",
        "05_client_markets",
        "06_client_share",
        "07_current_business",
    ):
        section = clean_brief.section(sid)
        assert section is not None and section.tables
        assert section.tables[0].num_cols == 2


def test_attendee_table_is_4_columns(clean_brief: Brief) -> None:
    section = clean_brief.section("11_meeting_with")
    assert section is not None
    assert section.tables[0].num_cols == 4


def test_bullets_detected_in_exec_messages(clean_brief: Brief) -> None:
    section = clean_brief.section("09_exec_messages")
    assert section is not None
    assert len(section.bullets) >= 3


def test_key_facts_split_subsections(clean_brief: Brief) -> None:
    a = clean_brief.section("08a_business_overview")
    b = clean_brief.section("08b_competition")
    c = clean_brief.section("08c_vertex_overview")
    d = clean_brief.section("08d_notable_changes")
    assert a is not None and a.present and "Acme" in a.raw_text
    assert b is not None and b.present and "PU" in b.raw_text
    assert c is not None and c.present and "cards in force" in c.raw_text.lower()
    assert d is not None and d.present and "CFO" in d.raw_text


def test_key_facts_two_column_table_label_value_rows() -> None:
    """Many templates use a 2-column table (label | body), sometimes with a blank Vertex label."""
    from io import BytesIO

    from docx import Document

    from core.docx_parser import parse_brief

    doc = Document()
    doc.add_paragraph("FirstBank Meeting Brief")
    doc.add_paragraph("What are the key facts about the client?")
    t = doc.add_table(rows=4, cols=2)
    t.rows[0].cells[0].text = "Business Overview"
    t.rows[0].cells[1].text = "FirstBank is a top-10 US retail bank with $280B in assets."
    t.rows[1].cells[0].text = "Competition"
    t.rows[1].cells[1].text = "PU holds 28% of FirstBank debit volume."
    t.rows[2].cells[0].text = ""
    t.rows[2].cells[1].text = (
        "Vertex has 12M cards in force with FirstBank across a $1.8B portfolio."
    )
    t.rows[3].cells[0].text = "Notable Changes"
    t.rows[3].cells[1].text = "New CFO appointed in Feb 2025."
    buf = BytesIO()
    doc.save(buf)
    brief = parse_brief(buf.getvalue())
    a = brief.section("08a_business_overview")
    b = brief.section("08b_competition")
    c = brief.section("08c_vertex_overview")
    d = brief.section("08d_notable_changes")
    assert a is not None and a.present and "FirstBank" in a.raw_text
    assert b is not None and b.present and "PU" in b.raw_text
    assert c is not None and c.present and "12M" in c.raw_text and "portfolio" in c.raw_text.lower()
    assert d is not None and d.present and "CFO" in d.raw_text


def test_notable_changes_empty_when_omitted_in_key_facts() -> None:
    from io import BytesIO

    from docx import Document

    from core.docx_parser import parse_brief

    doc = Document()
    doc.add_paragraph("Example Corp Meeting Brief")
    doc.add_paragraph("What are the key facts about the client?")
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    cell.text = ""
    cell.add_paragraph("Business Overview: Short.")
    cell.add_paragraph("Competition: PU and NP.")
    cell.add_paragraph("Vertex Overview: 1M cards. $1B PV.")
    buf = BytesIO()
    doc.save(buf)
    brief = parse_brief(buf.getvalue())
    nc = brief.section("08d_notable_changes")
    assert nc is not None
    assert nc.present is True
    assert nc.raw_text == ""


def test_notable_changes_optional_never_section_missing() -> None:
    from io import BytesIO

    from docx import Document

    from core.checkers.structure import check_section_presence
    from core.docx_parser import parse_brief

    doc = Document()
    doc.add_paragraph("Tiny Co Meeting Brief")
    buf = BytesIO()
    doc.save(buf)
    brief = parse_brief(buf.getvalue())
    nc = brief.section("08d_notable_changes")
    assert nc is not None and nc.present is True and nc.raw_text == ""
    findings = check_section_presence(brief)
    assert not any(f.section_id == "08d_notable_changes" for f in findings)


def test_stacked_intro_table_parses_sections_three_to_seven() -> None:
    """Templates may place 03–07 in one 2-column table (label | value per row)."""
    from io import BytesIO

    from docx import Document

    doc = Document()
    doc.add_paragraph("Example Corp Meeting Brief")
    tbl = doc.add_table(rows=5, cols=2)
    rows = [
        ("Client Name & Type", "Example Corp, issuer"),
        ("Meeting Objective", "Discuss renewal."),
        ("Client Market(s)", "US retail."),
        ("Client Share in Market", "Vertex share: 10%."),
        ("Our Current Business", "Active 5-yr deal."),
    ]
    for i, (left, right) in enumerate(rows):
        tbl.rows[i].cells[0].text = left
        tbl.rows[i].cells[1].text = right
    buf = BytesIO()
    doc.save(buf)
    brief = parse_brief(buf.getvalue())
    for sid in (
        "03_client_name_type",
        "04_meeting_objective",
        "05_client_markets",
        "06_client_share",
        "07_current_business",
    ):
        sec = brief.section(sid)
        assert sec is not None and sec.present
        assert sec.tables and sec.tables[0].num_cols == 2


def test_messy_stacked_table_key_facts_in_label_cell_and_pipe_topics() -> None:
    """Real briefs often merge Key Facts + body in one cell and pipe-merge later sections."""
    from io import BytesIO

    from docx import Document

    from core.checkers.structure import check_section_presence
    from core.docx_parser import parse_brief

    doc = Document()
    doc.add_paragraph("TestCo Annual Executive - Meeting Brief")
    tbl = doc.add_table(rows=8, cols=2)
    rows_data: list[tuple[str, str]] = [
        (
            "Client Name & Type:",
            "TestCo, a diversified financial institution.",
        ),
        ("Meeting Objective:", "Discuss relationship and future opportunities."),
        ("Client Market(s):", "US, Canada, Europe."),
        ("Client Share in Market:", "Vertex share about fifty percent."),
        ("Our Current Business:", "Seven-year partnership with pilots."),
        (
            "What are the key facts about the client?\n"
            "Business Overview: TestCo is large and diversified.\n"
            "| Competition: Works with Payments United and NewPay.\n"
            "| Vertex Overview: Vertex is embedded in payments.\n"
            "| [Only If Applicable]: Notable Changes: Recent leadership changes.",
            "What are the key facts about the client?\n"
            "Business Overview: TestCo is large and diversified.\n"
            "| Competition: Works with Payments United and NewPay.\n"
            "| Vertex Overview: Vertex is embedded in payments.\n"
            "| [Only If Applicable]: Notable Changes: Recent leadership changes.",
        ),
        (
            "What are the 3-5 messages you want the President to raise?\n"
            "First message about growth.\n"
            "| Second message about fraud.\n"
            "| Any issues or topics that the client will likely raise?\n"
            "Pricing and competitive concerns.",
            "What are the 3-5 messages you want the President to raise?\n"
            "First message about growth.\n"
            "| Second message about fraud.\n"
            "| Any issues or topics that the client will likely raise?\n"
            "Pricing and competitive concerns.",
        ),
        (
            "Who am I meeting with?\nWho is joining me from Vertex? Alex from sales.",
            "Who am I meeting with?\nWho is joining me from Vertex? Alex from sales.",
        ),
    ]
    for i, (left, right) in enumerate(rows_data):
        tbl.rows[i].cells[0].text = left
        tbl.rows[i].cells[1].text = right
    buf = BytesIO()
    doc.save(buf)
    brief = parse_brief(buf.getvalue())

    for sid in (
        "03_client_name_type",
        "04_meeting_objective",
        "05_client_markets",
        "06_client_share",
        "07_current_business",
        "08a_business_overview",
        "08b_competition",
        "08c_vertex_overview",
        "09_exec_messages",
        "10_client_topics",
        "11_meeting_with",
        "12_vertex_attendees",
    ):
        sec = brief.section(sid)
        assert sec is not None and sec.present, sid

    comp = brief.section("08b_competition")
    assert comp is not None and "Payments United" in comp.raw_text
    topics = brief.section("10_client_topics")
    assert topics is not None and "Pricing" in topics.raw_text
    exec_sec = brief.section("09_exec_messages")
    assert exec_sec is not None
    assert "First message about growth" in exec_sec.raw_text
    assert "President" in exec_sec.raw_text
    attendees = brief.section("12_vertex_attendees")
    assert attendees is not None and "Alex" in attendees.raw_text

    missing = check_section_presence(brief)
    assert not any("Client Name" in f.message for f in missing)


def test_attendee_bio_nested_table_cell_text_is_extracted() -> None:
    """Regression: Bio copy is often authored inside a nested Word table in the
    Bio cell. Paragraph-only reads miss it, so structure/semantic see an empty Bio."""
    from io import BytesIO

    from docx import Document
    from docx.shared import Inches

    from core.checkers.semantic import _format_meeting_attendee_table_for_semantic
    from tests.fixtures.builder import _png_bytes

    doc = Document()
    doc.add_paragraph("Stub Co Meeting Brief")
    outer = doc.add_table(rows=2, cols=2)
    outer.rows[0].cells[0].text = "Client Name & Type:"
    outer.rows[0].cells[1].text = "Stub Co, issuer"
    outer.rows[1].cells[0].text = (
        "Who am I meeting with?\n\nWho is joining me from Vertex? Alex Smith, SVP, NA"
    )
    right_cell = outer.rows[1].cells[1]
    right_cell.text = ""
    nested = right_cell.add_table(rows=2, cols=4)
    for i, h in enumerate(
        ("Name/Titles", "Photo", "Bio", "Previously met with Vertex exec?")
    ):
        nested.rows[0].cells[i].text = h
    nested.rows[1].cells[0].text = (
        "Gerri Kellman (GEH-ree)\nGeneral Counsel\nMs. Kellman\n"
        "gerri.kellman@example.com"
    )
    run = nested.rows[1].cells[1].paragraphs[0].add_run()
    run.add_picture(BytesIO(_png_bytes()), width=Inches(0.5))
    bio_cell = nested.rows[1].cells[2]
    bio_cell.text = ""
    inner_bio = bio_cell.add_table(rows=1, cols=1)
    inner_bio.rows[0].cells[0].text = (
        "Gerri Kellman serves as General Counsel and oversees legal and compliance."
    )
    nested.rows[1].cells[3].text = ""
    buf = BytesIO()
    doc.save(buf)
    brief = parse_brief(buf.getvalue())
    sec = brief.section("11_meeting_with")
    assert sec is not None and sec.present
    four_col = next(t for t in sec.tables if t.num_cols == 4 and t.num_rows >= 2)
    bio_extracted = four_col.rows[1][2].strip()
    assert "Gerri Kellman serves" in bio_extracted
    structured = _format_meeting_attendee_table_for_semantic(sec)
    assert "Bio:\nGerri Kellman serves" in structured
    assert "Bio:\n(empty)" not in structured


def test_missing_section_marked_not_present() -> None:
    from io import BytesIO

    from docx import Document

    doc = Document()
    doc.add_paragraph("Stub Co Meeting Brief")
    doc.add_paragraph("Client Name & Type")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Client"
    table.rows[0].cells[1].text = "Stub Co, issuer"
    buf = BytesIO()
    doc.save(buf)

    brief = parse_brief(buf.getvalue())
    assert brief.section("12_vertex_attendees") is not None
    assert brief.section("12_vertex_attendees").present is False
