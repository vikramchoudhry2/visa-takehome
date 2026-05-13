"""Programmatic .docx fixture builder.

Produces:
  * `build_clean_brief()` - a fully-compliant Vertex client briefing.
  * `build_dirty_brief()` - the same brief intentionally seeded with one
    of every rule violation, so a single integration test can exercise
    every checker.

Fixtures are built with python-docx and returned as `bytes` so tests can
feed them directly into `parse_brief`.
"""

from __future__ import annotations

import io
import struct
import zlib

from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Inches
from docx.table import Table

# A 1x1 PNG (transparent) so we have a real image to embed for the
# header icon and headshot fixtures. python-docx requires a real image
# stream when adding pictures.
_PNG_1X1 = (
    b"\x89PNG\r\n\x1a\n"
    + b"\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15\xc4\x89"
    + b"\x00\x00\x00\rIDATx\x9cc\xf8\x0f\x00\x00\x01\x01\x00\x05\x82\xfc\xa3W\x00\x00\x00\x00IEND\xaeB`\x82"
)


def _png_bytes() -> bytes:
    """Return a valid 1x1 transparent PNG."""
    sig = b"\x89PNG\r\n\x1a\n"

    def chunk(tag: bytes, data: bytes) -> bytes:
        crc = zlib.crc32(tag + data) & 0xFFFFFFFF
        return struct.pack(">I", len(data)) + tag + data + struct.pack(">I", crc)

    ihdr = chunk(
        b"IHDR",
        struct.pack(">IIBBBBB", 1, 1, 8, 6, 0, 0, 0),
    )
    raw = b"\x00" + b"\x00\x00\x00\x00"
    idat = chunk(b"IDAT", zlib.compress(raw))
    iend = chunk(b"IEND", b"")
    return sig + ihdr + idat + iend


def _add_2col_row(doc: DocxDocument, label: str, value: str) -> Table:
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = label
    table.rows[0].cells[1].text = value
    return table


def _add_header_icon(doc: DocxDocument, png_bytes: bytes) -> None:
    header = doc.sections[0].header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(png_bytes), width=Inches(0.6))


def _add_attendee_table(
    doc: DocxDocument,
    rows: list[tuple[str, bytes | None, str, str]],
) -> Table:
    """4-column attendee table: Name/Titles | Photo | Bio | Met before?"""
    table = doc.add_table(rows=1 + len(rows), cols=4)
    headers = ("Name/Titles", "Photo", "Bio", "Previously met with Vertex exec?")
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_idx, (name_block, photo, bio, met) in enumerate(rows, start=1):
        table.rows[r_idx].cells[0].text = name_block
        cell = table.rows[r_idx].cells[1]
        cell.text = ""
        if photo is not None:
            run = cell.paragraphs[0].add_run()
            run.add_picture(io.BytesIO(photo), width=Inches(0.5))
        else:
            cell.text = "(photo missing)"
        table.rows[r_idx].cells[2].text = bio
        table.rows[r_idx].cells[3].text = met
    return table


def build_clean_brief() -> bytes:
    """Build a fully-compliant client briefing for Acme Bank."""
    doc = Document()
    png = _png_bytes()
    _add_header_icon(doc, png)

    doc.add_paragraph("Acme Bank Meeting Brief")

    doc.add_paragraph("Client Name & Type")
    _add_2col_row(doc, "Client", "Acme Bank, issuer")

    doc.add_paragraph("Meeting Objective")
    _add_2col_row(doc, "Objective", "Renew core issuing partnership through 2030.")

    doc.add_paragraph("Client Market(s)")
    _add_2col_row(doc, "Markets", "US and Canada consumer credit.")

    doc.add_paragraph("Client Share in Market")
    _add_2col_row(
        doc,
        "Share",
        "Vertex share of Acme PV: 62% as of Mar 2025.\n"
        "Acme overall US issuer share by PV: 8% as of Mar 2025.",
    )

    doc.add_paragraph("Our Current Business")
    _add_2col_row(
        doc,
        "Deals",
        "Active 7-yr deal on consumer credit. Pending small-business credit RFP. "
        "Proposed commercial debit pilot for Q3.",
    )

    doc.add_paragraph("What are the key facts about the client?")
    key_facts = doc.add_table(rows=1, cols=1)
    cell = key_facts.rows[0].cells[0]
    cell.text = ""
    cell.add_paragraph(
        "Business Overview: Acme Bank is the 6th-largest US issuer by PV. "
        "Payments contribute roughly 30% of net revenue. "
        "Strong digital-first growth in the past two years."
    )
    cell.add_paragraph(
        "Competition: PU holds 22% network share at Acme; NP holds 14%. "
        "An RFP for commercial debit closes Apr 2025. "
        "Acme tracks competitive PV monthly."
    )
    cell.add_paragraph(
        "Vertex Overview: 18M cards in force at Acme. "
        "Portfolio size $42B in PV across consumer and small business. "
        "Renewal cycle began Jan 2025."
    )
    cell.add_paragraph(
        "Notable Changes: New CFO appointed Feb 2025. "
        "Reorganized payments team in Mar 2025."
    )

    doc.add_paragraph("What are the 3-5 messages you want the executive to raise?")
    messages = [
        "Reaffirm commitment to multi-year partnership; tie to roadmap for tokenization expansion.",
        "Position Vertex Insights pilot as differentiator vs PU bid; share early ROI data.",
        "Secure verbal alignment on small-business RFP scope before formal response in May 2025.",
        "Flag co-marketing investment for the holiday season; align on $25M shared spend.",
    ]
    for msg in messages:
        p = doc.add_paragraph(msg, style="List Bullet")
        _ = p

    doc.add_paragraph("Any issues or topics that the client will likely raise?")
    topics = [
        "Concern: interchange compression in commercial. Response: walk through Vertex protection clauses and roadmap.",
        "Concern: PU's aggressive pricing on debit. Response: highlight Vertex value-add services and incremental authorizations.",
        "Concern: cross-border fee transparency. Response: commit to a Q2 2025 review and dashboard delivery.",
    ]
    for t in topics:
        doc.add_paragraph(t, style="List Bullet")

    doc.add_paragraph("Who am I meeting with?")
    _add_attendee_table(
        doc,
        [
            (
                "Jane Doe (JAYN DOH)\nCEO\nMs. Doe\njane.doe@acme.com",
                png,
                "Jane has led Acme Bank since 2019. She spent 12 years at a major US issuer prior.",
                "Met President at Vertex Summit, Oct 2023, New York.",
            ),
            (
                "John Roe (JON ROH)\nCFO\nMr. Roe\njohn.roe@acme.com",
                png,
                "John joined as CFO in Feb 2025. Previously CFO at a regional fintech for six years.",
                "No prior meeting with Vertex executives.",
            ),
        ],
    )

    doc.add_paragraph("Who is joining me from Vertex?")
    doc.add_paragraph("Alex Smith, SVP, North America Issuing")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_dirty_brief() -> bytes:
    """Build a brief seeded with one violation per rule.

    Used by integration tests to verify every checker fires.
    """
    doc = Document()
    png = _png_bytes()
    _add_header_icon(doc, png)

    doc.add_paragraph("acme bank meeting brief")

    doc.add_paragraph("Client Name & Type")
    _add_2col_row(doc, "Client", "Acme Bank, issuer in the United States")

    doc.add_paragraph("Meeting Objective")
    _add_2col_row(doc, "Objective", "Renew the the core issuing partnership.")

    doc.add_paragraph("Client Market(s)")
    _add_2col_row(doc, "Markets", "United States and Canada consumer credit.")

    doc.add_paragraph("Client Share in Market")
    _add_2col_row(
        doc,
        "Share",
        "Vertex share of Acme PV: 62% as of January 2025.\n"
        "Acme US issuer share by PV: 8% as of January 25.",
    )

    doc.add_paragraph("Our Current Business")
    _add_2col_row(
        doc,
        "Deals",
        "Active 7 year deal on consumer credit worth $42 million.  Pending RFP.",
    )

    doc.add_paragraph("What are the key facts about the client?")
    key_facts = doc.add_table(rows=1, cols=1)
    cell = key_facts.rows[0].cells[0]
    cell.text = ""
    cell.add_paragraph(
        "Business Overview: Acme is a a bank that operates across many regions and segments. "
        "Payments contribute meaningfully to net revenue across consumer, commercial, and small business lines. "
        "Growth has accelerated thanks to a digital first strategy and renewed focus on co brand portfolios. "
        "The leadership team has tripled headcount in payments product over the last eighteen months, "
        "expanded into three new geographies, and launched a card refresh program."
    )
    cell.add_paragraph(
        "Competition: Payments United holds significant share at Acme across multiple product lines today. "
        "NewPay has been aggressively bidding on the small business segment and recently added a co brand partnership. "
        "Several other regional networks have also been pitching bundled offers to Acme over the past two quarters."
    )
    cell.add_paragraph(
        "Vertex Overview: Things are good and the relationship has been strong for many years. "
        "Portfolio is large across multiple product lines including consumer credit and commercial. "
        "More cards are coming online each quarter as the new digital onboarding flow ramps up across regions."
    )

    doc.add_paragraph("What are the 3-5 messages you want the executive to raise?")
    for i in range(6):
        doc.add_paragraph(
            f"We value our parntership with Acme and want to continue it message {i}.",
            style="List Bullet",
        )

    doc.add_paragraph("Any issues or topics that the client will likely raise?")
    for i in range(4):
        doc.add_paragraph(
            f"Generic concern {i} without a clear response plan.",
            style="List Bullet",
        )

    doc.add_paragraph("Who am I meeting with?")
    table = doc.add_table(rows=2, cols=4)
    table.rows[0].cells[0].text = "Name/Titles"
    table.rows[0].cells[1].text = "Photo"
    table.rows[0].cells[2].text = "Bio"
    table.rows[0].cells[3].text = "Previously met with Vertex exec?"
    table.rows[1].cells[0].text = "Jane Doe"
    table.rows[1].cells[1].text = ""
    table.rows[1].cells[2].text = "Short bio."
    table.rows[1].cells[3].text = "No"

    doc.add_paragraph("Who is joining me from Vertex?")
    doc.add_paragraph("Just Alex Smith")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
