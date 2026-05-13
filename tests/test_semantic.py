"""Semantic checker tests with a mocked LLM client."""

from __future__ import annotations

from core.checkers.semantic import (
    _section_text_for_semantic,
    check_semantic,
)
from core.llm.client import LLMFinding, LLMResult
from core.llm.prompts import SEMANTIC_RULES, SEMANTIC_RULES_BY_ID
from core.models import Brief, ParsedTable, Section


class FakeClient:
    def __init__(self, responses: dict[str, LLMResult] | None = None) -> None:
        self.responses = responses or {}
        self.calls: list[str] = []

    def evaluate(self, *, rule_id: str, system_prompt: str, user_prompt: str) -> LLMResult:
        self.calls.append(rule_id)
        return self.responses.get(rule_id, LLMResult(passed=True, findings=()))


def test_semantic_calls_every_rule_for_present_sections(clean_brief: Brief) -> None:
    fake = FakeClient()
    check_semantic(clean_brief, fake)
    expected = {r.rule_id for r in SEMANTIC_RULES}
    assert set(fake.calls) == expected


def test_semantic_emits_findings_with_correct_section_id(dirty_brief: Brief) -> None:
    fake = FakeClient(
        {
            "SEC09_PROACTIVE_MESSAGES": LLMResult(
                passed=False,
                findings=(LLMFinding(message="Generic talking point", evidence="value our parntership"),),
            ),
        }
    )
    findings = check_semantic(dirty_brief, fake)
    assert len(findings) == 1
    f = findings[0]
    assert f.section_id == "09_exec_messages"
    assert f.column == "section"
    assert f.rule_id == "SEC09_PROACTIVE_MESSAGES"
    assert "Generic" in f.message


def test_semantic_skips_missing_sections() -> None:
    from io import BytesIO

    from docx import Document

    from core.docx_parser import parse_brief

    doc = Document()
    doc.add_paragraph("Stub Co Meeting Brief")
    buf = BytesIO()
    doc.save(buf)

    brief = parse_brief(buf.getvalue())
    fake = FakeClient()
    findings = check_semantic(brief, fake)
    assert findings == []
    assert fake.calls == []


def test_semantic_skips_vertex_substance_when_value_cell_empty() -> None:
    """Regression: empty Vertex value cell still yields non-empty `raw_text`
    (the template label). The LLM rule must not run or it duplicates
    `VERTEX_ATTENDEE_MISSING` with a second near-identical bullet."""
    from io import BytesIO

    from docx import Document

    from core.docx_parser import parse_brief

    doc = Document()
    doc.add_paragraph("Test Co Meeting Brief")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Who is joining me from Vertex?"
    table.rows[0].cells[1].text = ""
    buf = BytesIO()
    doc.save(buf)
    brief = parse_brief(buf.getvalue())
    fake = FakeClient(
        {
            "SEC12_VERTEX_ATTENDEES_SUBSTANCE": LLMResult(
                passed=False,
                findings=(
                    LLMFinding(
                        message=(
                            'Section is empty. Must list Vertex attendees with name '
                            'and title/team or state "No other Vertex attendees".'
                        ),
                        evidence=None,
                    ),
                ),
            ),
        }
    )
    findings = check_semantic(brief, fake)
    assert "SEC12_VERTEX_ATTENDEES_SUBSTANCE" not in fake.calls
    assert findings == []


def test_semantic_skips_vertex_substance_when_line_format_invalid() -> None:
    """VERTEX_ATTENDEE_FORMAT already explains '[Name], [Title, Team]'; do not
    call the LLM for the same section (avoids duplicate bullets)."""
    from io import BytesIO

    from docx import Document

    from core.docx_parser import parse_brief

    doc = Document()
    doc.add_paragraph("Test Co Meeting Brief")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Who is joining me from Vertex?"
    table.rows[0].cells[1].text = "Just Alex Smith"
    buf = BytesIO()
    doc.save(buf)
    brief = parse_brief(buf.getvalue())
    fake = FakeClient()
    check_semantic(brief, fake)
    assert "SEC12_VERTEX_ATTENDEES_SUBSTANCE" not in fake.calls


def test_semantic_no_client_returns_empty(monkeypatch, dirty_brief: Brief) -> None:
    monkeypatch.delenv("ANTHROPIC_API_KEY", raising=False)
    findings = check_semantic(dirty_brief, client=None)
    assert findings == []


def test_semantic_section_11_uses_column_labeled_attendee_text() -> None:
    """Regression: flat raw_text stacks cells without headers; LLMs often
    misread Bio vs ``Previously met``. Section 11 semantic rules must see
    the same column alignment as structure checks."""
    rule_bio = SEMANTIC_RULES_BY_ID["SEC11_BIO_QUALITY"]
    attendee = ParsedTable(
        rows=(
            ("Name/Titles", "Photo", "Bio", "Previously met with Vertex exec?"),
            (
                "Jane Doe (JAYN DOH)\nVP\nPreferred: Ms. Doe\njane@acme.co",
                "",
                "SUBSTANTIVE_BIO_MARKER",
                "",
            ),
        ),
        has_images=False,
        cell_images=((False,) * 4, (False,) * 4),
    )
    section = Section(
        id="11_meeting_with",
        title="Who am I meeting with?",
        order=13,
        present=True,
        raw_text="SUBSTANTIVE_BIO_MARKER\n\nunlabeled tail",
        tables=(attendee,),
    )
    text = _section_text_for_semantic(section, rule_bio)
    assert "Bio:\nSUBSTANTIVE_BIO_MARKER" in text
    assert "Previously met with Vertex exec?:\n(empty)" in text


def test_sec11_previously_met_prompt_skips_empty_cell_duplication() -> None:
    """SEC11 must not repeat ATTENDEE_MET_EMPTY (structure already flags blank cells)."""
    rule = SEMANTIC_RULES_BY_ID["SEC11_PREVIOUSLY_MET"]
    assert "never emit a finding" in rule.system_prompt.lower()
    assert "duplicating" in rule.system_prompt.lower()
    assert "(empty)" in rule.user_prompt_template


def test_semantic_section_11_passes_structured_text_to_llm(clean_brief: Brief) -> None:
    captured: dict[str, str] = {}

    class CaptureClient:
        def evaluate(
            self, *, rule_id: str, system_prompt: str, user_prompt: str
        ) -> LLMResult:
            captured[rule_id] = user_prompt
            return LLMResult(passed=True, findings=())

    check_semantic(clean_brief, CaptureClient())
    bio_prompt = captured.get("SEC11_BIO_QUALITY", "")
    assert "Bio:\n" in bio_prompt
    assert "Attendee grid (column-aligned extract)" in bio_prompt
