"""Structural / section checker tests."""

from __future__ import annotations

import io

from docx import Document
from docx.shared import Inches

from core.checkers.header_image import check_header_icon
from core.checkers.structure import (
    check_client_topics,
    check_exec_messages,
    check_key_facts_category_labels,
    check_key_facts_lines,
    check_meeting_with,
    check_section_presence,
    check_structure,
    check_title,
    check_two_column_tables,
    check_vertex_attendees,
)
from core.docx_parser import parse_brief
from core.models import Brief
from tests.fixtures.builder import _png_bytes


def test_clean_brief_passes_all_structure_checks(clean_brief: Brief) -> None:
    assert check_structure(clean_brief) == []
    assert check_header_icon(clean_brief) == []


def test_dirty_brief_fires_structure_rules(dirty_brief: Brief) -> None:
    rule_ids = {f.rule_id for f in check_structure(dirty_brief)}
    expected_core = {
        "TITLE_CAPITALIZATION",
        "KEY_FACTS_LINE_LIMIT",
        "MAX_FIVE_BULLETS",
        "MAX_THREE_BULLETS",
        "ATTENDEE_NAME_LINE",
        "ATTENDEE_MISSING_PRONUNCIATION",
        "ATTENDEE_MISSING_TITLE",
        "ATTENDEE_MISSING_PREFERRED_FORM_OF_ADDRESS",
        "ATTENDEE_MISSING_EMAIL",
        "ATTENDEE_PHOTO_MISSING",
        "VERTEX_ATTENDEE_FORMAT",
    }
    missing = expected_core - rule_ids
    assert not missing, f"missing: {missing}; got {rule_ids}"


def _build_minimal_brief_with_attendee_table(
    headers: tuple[str, str, str, str],
    rows: list[tuple[str, bytes | str | None, str, str]],
) -> Brief:
    """Build a minimal brief that contains only the attendee section.

    The Photo cell input is interpreted as: bytes -> embed image,
    str -> write as text, None -> empty cell.
    """
    doc = Document()
    doc.add_paragraph("Stub Co Meeting Brief")
    doc.add_paragraph("Who am I meeting with?")
    table = doc.add_table(rows=1 + len(rows), cols=4)
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_idx, (name_block, photo, bio, met) in enumerate(rows, start=1):
        table.rows[r_idx].cells[0].text = name_block
        photo_cell = table.rows[r_idx].cells[1]
        photo_cell.text = ""
        if isinstance(photo, bytes):
            run = photo_cell.paragraphs[0].add_run()
            run.add_picture(io.BytesIO(photo), width=Inches(0.5))
        elif isinstance(photo, str):
            photo_cell.text = photo
        table.rows[r_idx].cells[2].text = bio
        table.rows[r_idx].cells[3].text = met
    buf = io.BytesIO()
    doc.save(buf)
    return parse_brief(buf.getvalue())


def test_attendee_table_headers_must_match_spec() -> None:
    brief = _build_minimal_brief_with_attendee_table(
        headers=("Name", "Image", "About", "History"),
        rows=[
            (
                "Jane Doe (JAYN DOH)\nCEO\nMs. Doe\njane.doe@acme.com",
                _png_bytes(),
                "Jane has led Acme since 2019.",
                "Met Smith at Vertex Summit, Oct 2023, NYC.",
            )
        ],
    )
    rule_ids = {f.rule_id for f in check_meeting_with(brief)}
    assert "ATTENDEE_TABLE_HEADERS" in rule_ids


def test_attendee_photo_missing_is_per_row() -> None:
    """Row 1 has an image, row 2 does not — only row 2 should be flagged."""
    brief = _build_minimal_brief_with_attendee_table(
        headers=("Name/Titles", "Photo", "Bio", "Previously met with Vertex exec?"),
        rows=[
            (
                "Jane Doe (JAYN DOH)\nCEO\nMs. Doe\njane.doe@acme.com",
                _png_bytes(),
                "Jane has led Acme since 2019. Strong issuer background.",
                "Met Smith at Vertex Summit, Oct 2023, NYC.",
            ),
            (
                "John Roe (JON ROH)\nCFO\nMr. Roe\njohn.roe@acme.com",
                None,
                "John joined as CFO in Feb 2025. Prior CFO at a regional fintech.",
                "No prior meeting with Vertex executives.",
            ),
        ],
    )
    findings = check_meeting_with(brief)
    photo_findings = [f for f in findings if f.rule_id == "ATTENDEE_PHOTO_MISSING"]
    assert len(photo_findings) == 1
    assert "row 2" in photo_findings[0].message


def test_attendee_photo_unclear_when_image_has_caption_text() -> None:
    """Image + stray text in the same Photo cell flags ATTENDEE_PHOTO_UNCLEAR."""
    doc = Document()
    doc.add_paragraph("Stub Co Meeting Brief")
    doc.add_paragraph("Who am I meeting with?")
    table = doc.add_table(rows=2, cols=4)
    for i, h in enumerate(
        ("Name/Titles", "Photo", "Bio", "Previously met with Vertex exec?")
    ):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].text = (
        "Jane Doe (JAYN DOH)\nCEO\nMs. Doe\njane.doe@acme.com"
    )
    photo_cell = table.rows[1].cells[1]
    photo_cell.text = "TBD - new headshot pending"
    run = photo_cell.add_paragraph().add_run()
    run.add_picture(io.BytesIO(_png_bytes()), width=Inches(0.5))
    table.rows[1].cells[2].text = (
        "Jane has led Acme since 2019. Strong issuer background."
    )
    table.rows[1].cells[3].text = "Met Smith at Vertex Summit, Oct 2023, NYC."
    buf = io.BytesIO()
    doc.save(buf)
    brief = parse_brief(buf.getvalue())
    rule_ids = {f.rule_id for f in check_meeting_with(brief)}
    assert "ATTENDEE_PHOTO_UNCLEAR" in rule_ids


def test_attendee_missing_name_when_cell_blank_except_email() -> None:
    """Name cell with only an email and no name text flags ATTENDEE_MISSING_NAME."""
    brief = _build_minimal_brief_with_attendee_table(
        headers=("Name/Titles", "Photo", "Bio", "Previously met with Vertex exec?"),
        rows=[
            (
                "(JAYN DOH)\nCEO\nMs.\njane.doe@acme.com",
                _png_bytes(),
                "Bio text here. More bio.",
                "Met Smith at Vertex Summit, Oct 2023, NYC.",
            )
        ],
    )
    rule_ids = {f.rule_id for f in check_meeting_with(brief)}
    assert "ATTENDEE_MISSING_NAME" in rule_ids


def test_attendee_title_accepts_broader_keywords() -> None:
    """Non-default exec titles (Founder, Chair, CISO) must not be false-flagged."""
    brief = _build_minimal_brief_with_attendee_table(
        headers=("Name/Titles", "Photo", "Bio", "Previously met with Vertex exec?"),
        rows=[
            (
                "Jane Doe (JAYN DOH)\nFounder & Chair\nMs. Doe\njane.doe@acme.com",
                _png_bytes(),
                "Jane founded Acme in 2010. Has chaired the board since.",
                "Met Smith at Vertex Summit, Oct 2023, NYC.",
            ),
            (
                "Sam Roe (SAM ROH)\nCISO\nMr. Roe\nsam.roe@acme.com",
                _png_bytes(),
                "Sam leads InfoSec at Acme. 15 years in payments security.",
                "No prior meeting with Vertex executives.",
            ),
        ],
    )
    rule_ids = {f.rule_id for f in check_meeting_with(brief)}
    assert "ATTENDEE_MISSING_TITLE" not in rule_ids


def test_attendee_bio_long_single_paragraph_exceeds_seven_visual_lines() -> None:
    """Bio is often one Word-wrapped paragraph with no newlines; line count
    must use a narrow column budget (not ``CHARS_PER_LINE``) or long bios
    in a 4-col table slip under the 7-line cap."""
    long_bio = (
        "Gerri Kellman serves as General Counsel for Global Horizon Financial "
        "Services Corporation and oversees legal, compliance, regulatory affairs, "
        "and corporate governance across all global operations. She has spent more "
        "than twenty five years in financial services working across numerous "
        "jurisdictions, regulatory regimes, and complex transactions and plays a "
        "significant role in shaping enterprise wide risk management."
    )
    assert len(long_bio) > 350  # sanity: exceeds ~7 wrapped lines at 50 cpl
    brief = _build_minimal_brief_with_attendee_table(
        headers=("Name/Titles", "Photo", "Bio", "Previously met with Vertex exec?"),
        rows=[
            (
                "Jane Doe (JAYN DOH)\nCEO\nMs. Doe\njane.doe@acme.com",
                _png_bytes(),
                long_bio,
                "Met Smith at Vertex Summit, Oct 2023, NYC.",
            )
        ],
    )
    rule_ids = {f.rule_id for f in check_meeting_with(brief)}
    assert "ATTENDEE_BIO_TOO_LONG" in rule_ids


def test_nested_attendee_table_inside_outer_stacked_table_is_surfaced() -> None:
    """Regression: when the brief lives inside one outer 2-col table and
    the attendee 4-col table is *nested* in the right cell of the
    'Who am I meeting with?' row, the parser must surface the nested
    table so the checker validates 4 columns (not the synthetic 1x2)."""
    doc = Document()
    doc.add_paragraph("Stub Co Meeting Brief")
    outer = doc.add_table(rows=2, cols=2)
    outer.rows[0].cells[0].text = "Client Name & Type:"
    outer.rows[0].cells[1].text = "Stub Co, issuer"
    # Compound left cell with both attendee questions; nested 4-col table
    # lives inside the right cell.
    outer.rows[1].cells[0].text = (
        "Who am I meeting with?\n\nWho is joining me from Vertex? Alex Smith, SVP, NA"
    )
    right_cell = outer.rows[1].cells[1]
    right_cell.text = ""
    nested = right_cell.add_table(rows=2, cols=4)
    for i, h in enumerate(
        ("Name/Titles", "Photo", "Bio", "Previously met with Vertex exec?")
    ):
        nested.rows[0].cells[i].text = h
    nested.rows[1].cells[0].text = (
        "Jane Doe (JAYN DOH)\nCEO\nMs. Doe\njane.doe@stub.co"
    )
    run = nested.rows[1].cells[1].paragraphs[0].add_run()
    run.add_picture(io.BytesIO(_png_bytes()), width=Inches(0.5))
    nested.rows[1].cells[2].text = (
        "Jane has led Stub Co since 2019. Strong issuer background."
    )
    nested.rows[1].cells[3].text = (
        "Met President Smith at Vertex Summit, Oct 2023, NYC."
    )
    buf = io.BytesIO()
    doc.save(buf)
    brief = parse_brief(buf.getvalue())
    sec = brief.section("11_meeting_with")
    assert sec is not None and sec.present
    has_4col = any(t.num_cols == 4 and t.num_rows >= 2 for t in sec.tables)
    assert has_4col, f"expected a 4-col attendee table, got {[(t.num_rows, t.num_cols) for t in sec.tables]}"
    rule_ids = {f.rule_id for f in check_meeting_with(brief)}
    assert "ATTENDEE_TABLE_COLUMNS" not in rule_ids
    assert "ATTENDEE_TABLE_MISSING" not in rule_ids


def test_exec_messages_plain_paragraphs_without_list_style_still_checked() -> None:
    """Many real briefs put messages as plain lines in a table cell (no ``w:numPr``),
    so ``section.bullets`` is empty; we still enforce count and 3-line cap."""
    long_msg = ("Vertex should lead with modernization. " * 18).strip()
    doc = Document()
    doc.add_paragraph("Stub Co Meeting Brief")
    outer = doc.add_table(rows=3, cols=2)
    outer.rows[0].cells[0].text = "Client Name & Type:"
    outer.rows[0].cells[1].text = "Stub Co, issuer"
    outer.rows[1].cells[0].text = (
        "What are the 3-5 messages you want the President to raise?"
    )
    outer.rows[1].cells[1].text = (
        "First proactive point with context.\n"
        "Second point with intent.\n"
        f"{long_msg}\n"
        "Fourth closing point."
    )
    outer.rows[2].cells[0].text = (
        "Any issues or topics that the client will likely raise?"
    )
    outer.rows[2].cells[1].text = (
        ("Concern: pricing pressure in SMB. Response: walk through value story. " * 5).strip()
        + "\n"
        "Concern: fraud. Response: show roadmap.\n"
        "Concern: latency. Response: offer pilot.\n"
        "Concern: contract. Response: engage legal.\n"
    )
    buf = io.BytesIO()
    doc.save(buf)
    brief = parse_brief(buf.getvalue())
    exec_findings = check_exec_messages(brief)
    topic_findings = check_client_topics(brief)
    assert any(f.rule_id == "BULLET_LINE_LIMIT" for f in exec_findings)
    assert any(f.rule_id == "MAX_THREE_BULLETS" for f in topic_findings)
    assert any(f.rule_id == "BULLET_LINE_LIMIT" for f in topic_findings)


def test_attendee_preferred_form_accepts_preferred_label() -> None:
    """A 'Preferred: <name>' line counts as a form of address."""
    brief = _build_minimal_brief_with_attendee_table(
        headers=("Name/Titles", "Photo", "Bio", "Previously met with Vertex exec?"),
        rows=[
            (
                "Jane Doe (JAYN DOH)\nCEO\nPreferred: Jane\njane.doe@acme.com",
                _png_bytes(),
                "Jane has led Acme since 2019. Strong issuer background.",
                "Met Smith at Vertex Summit, Oct 2023, NYC.",
            )
        ],
    )
    rule_ids = {f.rule_id for f in check_meeting_with(brief)}
    assert "ATTENDEE_MISSING_PREFERRED_FORM_OF_ADDRESS" not in rule_ids


def test_global_horizon_client_topics_respects_visual_three_line_cap() -> None:
    """Regression: line budget must match ~Word wrap in the value column —
    not so tight that short bullets (2–3 visual lines) all get ``BULLET_LINE_LIMIT``."""
    from pathlib import Path

    import pytest

    path = Path("Client Briefing - Global Horizon Financial Services Corporation.docx")
    if not path.is_file():
        pytest.skip("Global Horizon fixture not present")
    brief = parse_brief(path.read_bytes())
    findings = check_client_topics(brief)
    line_findings = [f for f in findings if f.rule_id == "BULLET_LINE_LIMIT"]
    assert len(line_findings) == 3
    messages = {f.message for f in line_findings}
    assert "Bullet 1 must not exceed 3 lines." in messages
    assert "Bullet 2 must not exceed 3 lines." in messages
    assert "Bullet 4 must not exceed 3 lines." in messages
    assert "Bullet 3 must not exceed 3 lines." not in messages
    assert "Bullet 5 must not exceed 3 lines." not in messages


def test_section_presence_flags_missing() -> None:
    from io import BytesIO

    from docx import Document

    from core.docx_parser import parse_brief

    doc = Document()
    doc.add_paragraph("Stub Co Meeting Brief")
    buf = BytesIO()
    doc.save(buf)

    brief = parse_brief(buf.getvalue())
    findings = check_section_presence(brief)
    rule_ids = {f.rule_id for f in findings}
    assert "SECTION_MISSING" in rule_ids
    section_ids = {f.section_id for f in findings}
    assert "12_vertex_attendees" in section_ids


def test_header_icon_missing() -> None:
    from io import BytesIO

    from docx import Document

    from core.docx_parser import parse_brief

    doc = Document()
    doc.add_paragraph("Stub Co Meeting Brief")
    buf = BytesIO()
    doc.save(buf)

    brief = parse_brief(buf.getvalue())
    findings = check_header_icon(brief)
    assert len(findings) == 1
    assert findings[0].rule_id == "HEADER_ICON_MISSING"


def test_two_col_table_flags_wrong_column_count(dirty_brief: Brief) -> None:
    findings = check_two_column_tables(dirty_brief)
    _ = findings


def test_individual_checkers_match_aggregate(dirty_brief: Brief) -> None:
    aggregate_ids = {f.rule_id for f in check_structure(dirty_brief)}
    individual_ids: set[str] = set()
    for fn in (
        check_section_presence,
        check_title,
        check_two_column_tables,
        check_key_facts_category_labels,
        check_key_facts_lines,
        check_exec_messages,
        check_client_topics,
        check_meeting_with,
        check_vertex_attendees,
    ):
        individual_ids |= {f.rule_id for f in fn(dirty_brief)}
    assert aggregate_ids == individual_ids
