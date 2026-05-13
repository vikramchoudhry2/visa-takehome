"""Formatting / grammar checker tests."""

from __future__ import annotations

from core.checkers.formatting import check_formatting
from core.models import Brief


def _rule_ids(brief: Brief) -> set[str]:
    return {f.rule_id for f in check_formatting(brief)}


def test_clean_brief_has_no_formatting_findings(clean_brief: Brief) -> None:
    findings = check_formatting(clean_brief)
    assert findings == []


def test_dirty_brief_fires_every_formatting_rule(dirty_brief: Brief) -> None:
    rule_ids = _rule_ids(dirty_brief)
    expected = {
        "ABBREVIATE_US",
        "DUPLICATE_WORD",
        "MONTH_ABBREVIATION",
        "FULL_YEAR",
        "DEAL_LENGTH",
        "DOLLAR_FORMAT",
        "SINGLE_SPACE_AFTER_PERIOD",
        "COMPETITOR_ABBREVIATION",
        "SPELLING",
    }
    missing = expected - rule_ids
    assert not missing, f"Did not fire: {missing}; got {rule_ids}"


def test_us_pronoun_not_flagged() -> None:
    from io import BytesIO

    from docx import Document

    from core.docx_parser import parse_brief

    doc = Document()
    doc.add_paragraph("Acme Meeting Brief")
    doc.add_paragraph("Client Name & Type")
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Client"
    t.rows[0].cells[1].text = "Acme will work with us in Apr 2025."
    buf = BytesIO()
    doc.save(buf)

    brief = parse_brief(buf.getvalue())
    findings = check_formatting(brief)
    assert all(f.rule_id != "ABBREVIATE_US" for f in findings)


def test_competitor_abbreviation_attribution(dirty_brief: Brief) -> None:
    findings = check_formatting(dirty_brief)
    competitor_findings = [f for f in findings if f.rule_id == "COMPETITOR_ABBREVIATION"]
    assert competitor_findings, "no competitor findings"
    assert all(f.section_id == "08b_competition" for f in competitor_findings)
